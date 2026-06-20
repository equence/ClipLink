from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path('/private/tmp/alampy-original.docx')
OUTPUT = ROOT / 'ClipLink实验报告.docx'
CLIENT_IMAGE = Path('/var/folders/br/r2mssjl511gbvw389s6_l1lw0000gn/T/codex-clipboard-bacc17cf-3df8-48c3-89ea-a7338330db70.png')
SERVER_IMAGE = Path('/var/folders/br/r2mssjl511gbvw389s6_l1lw0000gn/T/codex-clipboard-1d255a7f-a5f6-45f2-9329-327469f64f4a.png')


def set_font(run, size=10.5, bold=False):
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold


def add_before(reference, text='', size=10.5, bold=False, heading=False):
    paragraph = reference.insert_paragraph_before()
    if heading:
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
    else:
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.35
    run = paragraph.add_run(text)
    set_font(run, size, bold)
    return paragraph


def remove_third_section(doc):
    paragraphs = list(doc.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith('实验三'))
    end = next(i for i, p in enumerate(paragraphs[start + 1:], start + 1) if p.text.strip().startswith('六、参考书'))
    for paragraph in paragraphs[start:end]:
        paragraph._element.getparent().remove(paragraph._element)
    return next(p for p in doc.paragraphs if p.text.strip().startswith('六、参考书'))


def add_image_before(reference, image, caption):
    paragraph = reference.insert_paragraph_before()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Cm(15.5))
    caption_p = reference.insert_paragraph_before()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_p.add_run(caption)
    set_font(run, 9)
    return caption_p


def main():
    doc = Document(SOURCE)
    reference = remove_third_section(doc)

    title = add_before(reference, '实验三  网络应用程序设计', 16, True, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.page_break_before = True

    add_before(reference, '实验目的及内容', 12, True, True)
    add_before(reference, '通过设计并实现 ClipLink 跨平台局域网剪贴板同步工具，掌握基于 TCP/IP 协议栈的 Socket 程序设计方法；理解客户机—服务器模型、TCP 长连接、消息封装、并发转发及图形界面事件驱动编程。')

    add_before(reference, '实验要求', 12, True, True)
    add_before(reference, 'ClipLink 采用客户机—服务器结构：服务端负责监听端口、维护连接并转发同步消息；客户端监听本机纯文本剪贴板变化，在局域网内将内容发送到服务端，并接收其他客户端的同步内容。客户端应提供清晰的图形界面，显示连接状态、同步历史和文本详情。')

    add_before(reference, '实验环境', 12, True, True)
    for item in ['1、微型计算机（macOS / Windows / Linux）', '2、C++17、CMake 3.21+', '3、Qt 6（Core、Network、Widgets、Test 模块）', '4、Visual Studio Code 或其他 C++ 开发环境']:
        add_before(reference, item)

    add_before(reference, '系统设计', 12, True, True)
    add_before(reference, '1、通信结构：ClipLink Server 使用 QTcpServer 监听默认端口 45454；多个 ClipLink Client 使用 QTcpSocket 建立 TCP 长连接。服务端校验消息后向除发送端外的在线客户端广播。')
    add_before(reference, '2、消息协议：应用层采用“4 字节大端长度 + UTF-8 JSON”的帧格式，JSON 包含 type、id、deviceId、deviceName 和 text 字段。长度前缀可解决 TCP 粘包与半包问题，单条消息限制为 16 KiB。')
    add_before(reference, '3、循环抑制：每条同步消息带 UUID；客户端使用最近消息 ID 缓存，并记录远端写入的剪贴板文本，避免接收后的剪贴板变化再次上传形成循环。')
    add_before(reference, '4、界面层：Qt Widgets 实现桌面客户端。主窗口包含连接状态、设置、历史列表、文本详情和手动发送区域；界面跟随系统浅色/深色模式。')

    add_before(reference, '实验及设计步骤', 12, True, True)
    for item in [
        '1、使用 CMake 配置工程，链接 Qt6::Core、Qt6::Network、Qt6::Widgets 和 Qt6::Test。',
        '2、实现协议编解码模块，完成 JSON 消息序列化、长度校验及非法帧拒绝。',
        '3、实现 RelayServer：接受多个 TCP 客户端连接，缓存不完整帧并将合法剪贴板消息转发给其他客户端。',
        '4、实现客户端 NetworkClient 与 ClipboardWatcher：监听系统剪贴板、发送本地文本、接收远端文本，并在断线时有限次数重连。',
        '5、实现图形界面并启动服务端。客户端点击“连接服务器”，填写服务端局域网地址和端口 45454 后即可同步纯文本。',
        '6、使用 QtTest 对协议编解码、消息去重、服务端转发、客户端连接、剪贴板监听和主窗口历史记录进行自动化测试。'
    ]:
        add_before(reference, item)

    add_before(reference, '实验报告', 12, True, True)
    add_before(reference, 'ClipLink 是一个使用 C++17 与 Qt 6 实现的跨平台局域网剪贴板同步工具，源代码仓库地址为：https://github.com/equence/ClipLink 。项目由无界面服务端和桌面 GUI 客户端组成。服务端负责 TCP 消息转发；客户端负责剪贴板监听、网络通信、历史显示以及手动发送。')
    add_before(reference, '运行时先执行 cliplink_server --port 45454 启动服务端，再运行 cliplink_client 启动客户端。客户端的同步历史列表会显示本地复制或远端接收的文本；选择一条记录可在右侧查看完整内容。')

    add_before(reference, '运行效果如下：', 12, True, True)
    add_image_before(reference, CLIENT_IMAGE, '图 1  ClipLink 客户端主界面与本地剪贴板历史记录')
    add_image_before(reference, SERVER_IMAGE, '图 2  ClipLink TCP 转发服务端监听 45454 端口')

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
